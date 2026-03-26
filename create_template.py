import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("python-docx NOT INSTALLED")
    sys.exit(1)

def create_proposal():
    doc = Document()

    # Section 1: Title Page
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # Set margins
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add Logo
    logo_path = r'c:\Brains WEB\public\learning-brains-logo-transparent-cropped.png'
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(3))
    
    # Add Space
    for _ in range(3):
        doc.add_paragraph()

    # Integrated On-the-job Learning Systems for Industrial Reskilling
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Integrated On-the-job Learning Systems for\nIndustrial Reskilling")
    run.font.size = Pt(18)
    run.font.name = 'Segoe UI'
    run.font.color.rgb = RGBColor(0x3B, 0x59, 0x98) # A blue-ish color

    # Code
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2025-1-ES01-KA220-VET-000351934")
    run.font.size = Pt(12)
    run.font.name = 'Segoe UI'
    
    # Add Space
    for _ in range(4):
        doc.add_paragraph()

    # TITLE
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TITLE")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Segoe UI'
    run.font.color.rgb = RGBColor(0x3B, 0x59, 0x98)

    # Footer for First Page
    footer = section.first_page_footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Layout the footer with a table to have text and image side by side
    table = footer.add_table(1, 2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Left cell: Text
    cell_text = table.rows[0].cells[0]
    p_footer = cell_text.paragraphs[0]
    text = (
        "The project is co-funded by the Erasmus+ Programme of the European Union. "
        "The content of this document represents the views of the author only and is "
        "his/her sole responsibility. The European Commission does not accept any "
        "responsibility for use that may be made of the information it contains."
    )
    run_footer = p_footer.add_run(text)
    run_footer.font.size = Pt(7)
    run_footer.font.italic = True
    
    # Right cell: EU Logo
    cell_logo = table.rows[0].cells[1]
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    eu_logo_path = r'c:\Brains WEB\public\eu-emblem.png'
    if os.path.exists(eu_logo_path):
        run_eu = p_logo.add_run()
        run_eu.add_picture(eu_logo_path, width=Inches(1.5))

    # Save
    output_path = r'c:\Brains WEB\Propuesta_Template_LearningBrains.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_proposal()
